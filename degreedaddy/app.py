import os
import io
from flask import Flask, request, render_template_string, jsonify, send_file
from openai import OpenAI
from docx import Document
from fpdf import FPDF
from dotenv import load_dotenv
import nest_asyncio

# Allow Flask to run inside Jupyter Notebook
nest_asyncio.apply()

# Load environment variables from .env file
load_dotenv()

# Initialize OpenAI client with API key from environment
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Flask app
app = Flask(__name__)

# Store conversation history
conversation_history = []

# ---------------- HTML TEMPLATE ----------------
html_template = """
<!DOCTYPE html>
<html>
<head>
<title>Interactive Course Tutor</title>
<style>
body { font-family: Arial, sans-serif; margin: 30px; line-height: 1.6; font-size: 16px; }
textarea { width: 100%; font-size: 16px; padding: 10px; margin-top: 5px; margin-bottom: 10px; border-radius: 5px; border: 1px solid #ccc; }
#output, #answer { margin-top: 20px; padding: 15px; border: 1px solid #ccc; white-space: pre-wrap; height: 300px; overflow-y: auto; font-size: 16px; line-height: 1.5; border-radius: 5px; background-color: #f9f9f9; }
button { margin-top: 10px; padding: 10px 20px; font-size: 16px; border-radius: 5px; cursor: pointer; border: none; color: white; }
button#start { background-color: #28a745; }
button#ask { background-color: #007bff; }
button#download-word { background-color: #17a2b8; }
button#download-pdf { background-color: #6c757d; }
h1, h3 { margin-bottom: 10px; }
</style>
</head>
<body>
<h1>Interactive Course Tutor</h1>

<h3>Step 1: Enter Syllabus</h3>
<textarea id="syllabus" rows="5" placeholder="Enter syllabus here..."></textarea><br>
<button id="start" onclick="startCourse()">Start Course</button>
<div id="output"></div>

<h3>Step 2: Ask Questions</h3>
<textarea id="question" rows="3" placeholder="Type your question here..."></textarea><br>
<button id="ask" onclick="askQuestion()">Ask</button>
<div id="answer"></div>

<h3>Step 3: Download Conversation</h3>
<button id="download-word" onclick="download('word')">Download Word</button>
<button id="download-pdf" onclick="download('pdf')">Download PDF</button>

<script>
async function startCourse() {
    const syllabus = document.getElementById("syllabus").value;
    document.getElementById("output").innerText = "Generating explanation...";
    const res = await fetch("/start", {
        method: "POST",
        headers: {"Content-Type":"application/json"},
        body: JSON.stringify({syllabus})
    });
    const data = await res.json();
    document.getElementById("output").innerText = data.explanation;
}

async function askQuestion() {
    const question = document.getElementById("question").value;
    const res = await fetch("/ask", {
        method: "POST",
        headers: {"Content-Type":"application/json"},
        body: JSON.stringify({question})
    });
    const data = await res.json();
    document.getElementById("answer").innerText = data.answer;
}

function download(format) {
    window.location.href = `/download?format=${format}`;
}
</script>
</body>
</html>
"""

# ---------------- ROUTES ----------------

@app.route('/')
def index():
    return render_template_string(html_template)

@app.route('/start', methods=['POST'])
def start_course():
    global conversation_history
    conversation_history = []  # reset
    syllabus = request.json['syllabus']

    prompt = f"""You are a helpful tutor for beginners.
Explain the prerequisites and each module with simple examples for this syllabus:
{syllabus}
"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.7,
        max_tokens=2000
    )

    explanation = response.choices[0].message.content
    conversation_history.append(("Course Explanation", explanation))
    return jsonify({"explanation": explanation})

@app.route('/ask', methods=['POST'])
def ask_question():
    global conversation_history
    question = request.json['question']
    conversation_history.append(("User Question", question))

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": question}],
        temperature=0.7,
        max_tokens=800
    )

    answer = response.choices[0].message.content
    conversation_history.append((question, answer))
    return jsonify({"answer": answer})

@app.route('/download')
def download():
    format_choice = request.args.get('format', 'pdf').lower()

    if format_choice == 'word':
        doc = Document()
        doc.add_heading("Interactive Course Tutor", 0)
        for q, a in conversation_history:
            doc.add_page_break()
            doc.add_heading(q, level=1)
            doc.add_paragraph(a)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name="conversation.docx")

    else:  # PDF
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, "Interactive Course Tutor", ln=True, align='C')
        pdf.ln(10)
        pdf.set_font("Arial", '', 12)
        for q, a in conversation_history:
            pdf.add_page()
            pdf.set_font("Arial", 'B', 14)
            pdf.multi_cell(0, 8, q)
            pdf.ln(5)
            pdf.set_font("Arial", '', 12)
            pdf.multi_cell(0, 8, a)
        buffer = io.BytesIO()
        pdf.output(buffer)
        buffer.seek(0)
        return send_file(buffer, as_attachment=True, download_name="conversation.pdf")

# ---------------- MAIN ----------------
if __name__ == '__main__':
    app.run(host="0.0.0.0", port=int(os.getenv("PORT", 5000)))
